import io
import json
import logging
from typing import Any, Dict, List, Optional

import docx
import markdown
import PyPDF2
from core.config import settings
from db.models import Document, DocumentChunk
from fastapi import UploadFile
from schemas.schemas import DocumentCreate, DocumentResponse
from services.vector_store import vector_store
from sqlalchemy.orm import Session

logger = logging.getLogger(__name__)


class DocumentService:
    """
    Service for document processing and management
    Handles chunking, embedding, and database operations
    """

    def __init__(self):
        pass

    def chunk_text(
        self, text: str, chunk_size: int = None, overlap: int = None
    ) -> List[str]:
        """
        Split text into overlapping chunks

        Args:
            text: Input text to chunk
            chunk_size: Size of each chunk in characters
            overlap: Number of characters to overlap

        Returns:
            List of text chunks
        """
        chunk_size = chunk_size or settings.CHUNK_SIZE
        overlap = overlap or settings.CHUNK_OVERLAP

        if len(text) <= chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size

            # Try to break at sentence boundary
            if end < len(text):
                last_period = text.rfind(".", start, end)
                if last_period > start + (chunk_size // 2):
                    end = last_period + 1

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = end - overlap

        logger.debug(f"Split text into {len(chunks)} chunks")
        return chunks

    async def create_document(
        self, db: Session, document_data: DocumentCreate
    ) -> DocumentResponse:
        """
        Create a new document with chunks and embeddings

        Args:
            db: Database session
            document_data: Document creation data

        Returns:
            Created document response
        """
        try:
            logger.info(f"Creating document: '{document_data.title}'")

            # Create document record
            db_document = Document(
                title=document_data.title,
                content=document_data.content,
                source=document_data.source,
                category=document_data.category,
                metadata_json=(
                    json.dumps(document_data.metadata)
                    if document_data.metadata
                    else None
                ),
            )

            db.add(db_document)
            db.flush()

            # Chunk the document
            chunks = self.chunk_text(document_data.content)
            db_document.chunk_count = len(chunks)

            # Create chunk records and prepare for vector storage
            chunk_texts = []
            chunk_metadatas = []
            chunk_ids = []

            for i, chunk_text in enumerate(chunks):
                # Create chunk record
                chunk_id = f"doc_{db_document.id}_chunk_{i}"

                db_chunk = DocumentChunk(
                    document_id=db_document.id,
                    chunk_text=chunk_text,
                    chunk_index=i,
                    vector_id=chunk_id,
                )
                db.add(db_chunk)

                # Prepare for vector storage
                chunk_texts.append(chunk_text)
                chunk_ids.append(chunk_id)

                # Build metadata with user_id and conversation_id for filtering
                metadata = {
                    "document_id": db_document.id,
                    "title": document_data.title,
                    "chunk_index": i,
                    "category": document_data.category or "general",
                    "source": document_data.source or "",
                }

                # Add user_id and conversation_id if they exist
                if db_document.user_id is not None:
                    metadata["user_id"] = db_document.user_id
                if db_document.conversation_id is not None:
                    metadata["conversation_id"] = db_document.conversation_id

                chunk_metadatas.append(metadata)

            # Add to vector store
            if chunk_texts:
                await vector_store.add_documents(
                    texts=chunk_texts, metadatas=chunk_metadatas, ids=chunk_ids
                )

            db.commit()
            db.refresh(db_document)

            logger.info(
                f"Successfully created document {db_document.id} with {len(chunks)} chunks"
            )

            return DocumentResponse.from_orm_model(db_document)

        except Exception as e:
            db.rollback()
            logger.error(f"Error creating document: {str(e)}", exc_info=True)
            raise

    def get_document(self, db: Session, document_id: int) -> Optional[DocumentResponse]:
        """
        Get a document by ID

        Args:
            db: Database session
            document_id: Document ID

        Returns:
            Document response or None
        """
        try:
            db_document = db.query(Document).filter(Document.id == document_id).first()
            if db_document:
                return DocumentResponse.from_orm_model(db_document)
            return None
        except Exception as e:
            logger.error(f"Error getting document {document_id}: {str(e)}")
            return None

    def list_documents(
        self,
        db: Session,
        skip: int = 0,
        limit: int = 100,
        category: Optional[str] = None,
    ) -> List[DocumentResponse]:
        """
        List documents with optional filtering

        Args:
            db: Database session
            skip: Number of records to skip
            limit: Maximum number of records to return
            category: Optional category filter

        Returns:
            List of document responses
        """
        try:
            query = db.query(Document)

            if category:
                query = query.filter(Document.category == category)

            documents = query.offset(skip).limit(limit).all()
            return [DocumentResponse.from_orm_model(doc) for doc in documents]

        except Exception as e:
            logger.error(f"Error listing documents: {str(e)}")
            return []

    async def delete_document(self, db: Session, document_id: int) -> bool:
        """
        Delete a document and its chunks

        Args:
            db: Database session
            document_id: Document ID to delete

        Returns:
            Success status
        """
        try:
            logger.info(f"Deleting document {document_id}")

            # Get document to find chunk IDs
            db_document = db.query(Document).filter(Document.id == document_id).first()
            if not db_document:
                logger.warning(f"Document {document_id} not found")
                return False

            # Get chunk IDs for vector store deletion
            chunks = (
                db.query(DocumentChunk)
                .filter(DocumentChunk.document_id == document_id)
                .all()
            )
            chunk_ids = [chunk.vector_id for chunk in chunks if chunk.vector_id]

            # Delete from vector store
            if chunk_ids:
                await vector_store.delete_documents(chunk_ids)

            # Delete from database (cascades to chunks)
            db.delete(db_document)
            db.commit()

            logger.info(f"Successfully deleted document {document_id}")
            return True

        except Exception as e:
            db.rollback()
            logger.error(
                f"Error deleting document {document_id}: {str(e)}", exc_info=True
            )
            return False

    async def bulk_create_documents(
        self, db: Session, documents: List[DocumentCreate]
    ) -> Dict[str, Any]:
        """
        Create multiple documents in batch

        Args:
            db: Database session
            documents: List of documents to create

        Returns:
            Dictionary with success/failure counts
        """
        success_count = 0
        failed_count = 0
        errors = []

        for i, doc_data in enumerate(documents):
            try:
                await self.create_document(db, doc_data)
                success_count += 1
            except Exception as e:
                failed_count += 1
                error_msg = f"Document {i} ('{doc_data.title}'): {str(e)}"
                errors.append(error_msg)
                logger.error(error_msg)

        return {
            "success_count": success_count,
            "failed_count": failed_count,
            "total": len(documents),
            "errors": errors,
        }

    def get_document_stats(self, db: Session) -> Dict[str, Any]:
        """
        Get statistics about documents

        Args:
            db: Database session

        Returns:
            Dictionary with statistics
        """
        try:
            total_docs = db.query(Document).count()
            total_chunks = db.query(DocumentChunk).count()

            categories = db.query(Document.category).distinct().all()
            category_list = [cat[0] for cat in categories if cat[0]]

            return {
                "total_documents": total_docs,
                "total_chunks": total_chunks,
                "categories": category_list,
                "avg_chunks_per_doc": (
                    round(total_chunks / total_docs, 2) if total_docs > 0 else 0
                ),
            }
        except Exception as e:
            logger.error(f"Error getting document stats: {str(e)}")
            return {}

    async def extract_text_from_file(self, file: UploadFile) -> str:
        """
        Extract text content from uploaded file

        Args:
            file: Uploaded file object

        Returns:
            Extracted text content

        Raises:
            ValueError: If file type is not supported
        """
        filename = file.filename.lower()
        content_type = file.content_type or ""

        try:
            # Read file content
            content = await file.read()

            # PDF files
            if filename.endswith(".pdf") or "pdf" in content_type:
                logger.info(f"Extracting text from PDF: {filename}")
                pdf_file = io.BytesIO(content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                text_parts = []
                for page in pdf_reader.pages:
                    text_parts.append(page.extract_text())
                return "\n\n".join(text_parts)

            # DOCX files
            elif (
                filename.endswith(".docx")
                or "word" in content_type
                or "document" in content_type
            ):
                logger.info(f"Extracting text from DOCX: {filename}")
                docx_file = io.BytesIO(content)
                doc = docx.Document(docx_file)
                text_parts = [paragraph.text for paragraph in doc.paragraphs]
                return "\n\n".join(text_parts)

            # Markdown files
            elif filename.endswith(".md") or filename.endswith(".markdown"):
                logger.info(f"Extracting text from Markdown: {filename}")
                text = content.decode("utf-8")
                # Convert markdown to plain text (removes formatting)
                html = markdown.markdown(text)
                # Simple HTML tag removal
                import re

                text = re.sub("<[^<]+?>", "", html)
                return text

            # Plain text files
            elif filename.endswith(".txt") or "text" in content_type:
                logger.info(f"Reading plain text file: {filename}")
                return content.decode("utf-8")

            else:
                raise ValueError(
                    f"Unsupported file type: {filename}. Supported types: PDF, DOCX, TXT, MD"
                )

        except Exception as e:
            logger.error(f"Error extracting text from file: {str(e)}", exc_info=True)
            raise ValueError(f"Failed to extract text from file: {str(e)}")

    async def upload_and_process_file(
        self,
        db: Session,
        file: UploadFile,
        user_id: int,
        conversation_id: Optional[int] = None,
        category: Optional[str] = None,
    ) -> DocumentResponse:
        """
        Upload a file, extract text, and process into vector store

        Args:
            db: Database session
            file: Uploaded file
            user_id: User ID who owns the document
            category: Optional category for the document

        Returns:
            Created document response
        """
        try:
            # Validate file size
            max_size = settings.MAX_FILE_SIZE_MB * 1024 * 1024
            file.file.seek(0, 2)  # Seek to end
            file_size = file.file.tell()
            file.file.seek(0)  # Reset to beginning

            if file_size > max_size:
                raise ValueError(
                    f"File size ({file_size / 1024 / 1024:.1f} MB) exceeds maximum allowed size ({settings.MAX_FILE_SIZE_MB} MB)"
                )

            logger.info(
                f"Processing uploaded file: {file.filename} ({file_size / 1024:.1f} KB)"
            )

            # Extract text from file
            content = await self.extract_text_from_file(file)

            if not content or len(content.strip()) < 10:
                raise ValueError("Extracted content is too short or empty")

            # Create document data
            document_data = DocumentCreate(
                title=file.filename,
                content=content,
                source=f"user_upload:{file.filename}",
                category=category or "user_document",
                metadata={"uploaded_by": user_id, "filename": file.filename},
            )

            # Create document with user_id
            logger.info(f"Creating document record for user {user_id}")

            # Create document record
            db_document = Document(
                user_id=user_id,
                conversation_id=conversation_id,
                title=document_data.title,
                content=document_data.content,
                source=document_data.source,
                category=document_data.category,
                metadata_json=(
                    json.dumps(document_data.metadata)
                    if document_data.metadata
                    else None
                ),
            )

            db.add(db_document)
            db.flush()

            # Chunk the document
            chunks = self.chunk_text(document_data.content)
            db_document.chunk_count = len(chunks)

            # Create chunk records and prepare for vector storage
            chunk_texts = []
            chunk_metadatas = []
            chunk_ids = []

            for i, chunk_text in enumerate(chunks):
                vector_id = f"doc_{db_document.id}_chunk_{i}"

                db_chunk = DocumentChunk(
                    document_id=db_document.id,
                    chunk_text=chunk_text,
                    chunk_index=i,
                    vector_id=vector_id,
                )
                db.add(db_chunk)

                chunk_texts.append(chunk_text)
                chunk_ids.append(vector_id)
                chunk_metadatas.append(
                    {
                        "document_id": db_document.id,
                        "title": db_document.title,
                        "source": db_document.source,
                        "category": db_document.category,
                        "chunk_index": i,
                        "user_id": user_id,
                        "conversation_id": conversation_id,
                    }
                )

            # Add to vector store
            await vector_store.add_documents(chunk_texts, chunk_metadatas, chunk_ids)

            # Commit transaction
            db.commit()
            db.refresh(db_document)

            logger.info(
                f"Successfully processed file: {file.filename} (ID: {db_document.id})"
            )

            return DocumentResponse.from_orm_model(db_document)

        except ValueError:
            raise
        except Exception as e:
            db.rollback()
            logger.error(f"Error processing uploaded file: {str(e)}", exc_info=True)
            raise

    def list_user_documents(
        self,
        db: Session,
        user_id: int,
        skip: int = 0,
        limit: int = 100,
        category: Optional[str] = None,
        session_id: Optional[str] = None,
    ) -> List[DocumentResponse]:
        """
        List documents owned by a specific user
        Optionally filter by session_id to get conversation-specific documents

        Args:
            db: Database session
            user_id: User ID
            skip: Number of records to skip
            limit: Maximum number of records to return
            category: Optional category filter
            session_id: Optional session ID to filter by conversation

        Returns:
            List of document responses
        """
        try:
            query = db.query(Document).filter(Document.user_id == user_id)

            if category:
                query = query.filter(Document.category == category)

            # Filter by conversation if session_id is provided
            if session_id:
                from db.models import Conversation

                conversation = (
                    db.query(Conversation)
                    .filter(
                        Conversation.session_id == session_id,
                        Conversation.user_id == user_id,
                    )
                    .first()
                )
                if conversation:
                    query = query.filter(Document.conversation_id == conversation.id)
                else:
                    # No conversation found, return empty list
                    return []

            documents = (
                query.order_by(Document.created_at.desc())
                .offset(skip)
                .limit(limit)
                .all()
            )

            return [DocumentResponse.from_orm_model(doc) for doc in documents]

        except Exception as e:
            logger.error(f"Error listing user documents: {str(e)}")
            return []

    async def delete_user_document(
        self, db: Session, document_id: int, user_id: int
    ) -> bool:
        """
        Delete a document owned by a specific user

        Args:
            db: Database session
            document_id: Document ID to delete
            user_id: User ID for ownership verification

        Returns:
            Success status
        """
        try:
            logger.info(f"User {user_id} deleting document {document_id}")

            # Get document and verify ownership
            db_document = (
                db.query(Document)
                .filter(Document.id == document_id, Document.user_id == user_id)
                .first()
            )

            if not db_document:
                logger.warning(
                    f"Document {document_id} not found or not owned by user {user_id}"
                )
                return False

            # Get chunk IDs for vector store deletion
            chunks = (
                db.query(DocumentChunk)
                .filter(DocumentChunk.document_id == document_id)
                .all()
            )
            chunk_ids = [chunk.vector_id for chunk in chunks if chunk.vector_id]

            # Delete from vector store
            if chunk_ids:
                await vector_store.delete_documents(chunk_ids)

            # Delete from database (cascades to chunks)
            db.delete(db_document)
            db.commit()

            logger.info(f"Successfully deleted document {document_id}")
            return True

        except Exception as e:
            db.rollback()
            logger.error(
                f"Error deleting user document {document_id}: {str(e)}", exc_info=True
            )
            return False


# Global document service instance
document_service = DocumentService()
